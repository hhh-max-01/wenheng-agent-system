from __future__ import annotations

import base64
import binascii
import csv
import io
import json
import os
import re
import shutil
import struct
import subprocess
import tempfile
import urllib.error
import urllib.request
import zipfile
import zlib
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any

from docx import Document
from docx.exceptions import InvalidSpanError
from lxml import etree
from pypdf import PdfReader

try:
    from office_oxide import extract_text as oxide_extract_text
    from office_oxide import to_markdown as oxide_to_markdown
except ImportError:  # Local development can still use the LibreOffice fallback.
    oxide_extract_text = None
    oxide_to_markdown = None


ROOT = Path(__file__).resolve().parent
WEB_ROOT = ROOT / "web"


def load_local_env() -> None:
    env_path = ROOT / ".env"
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8-sig").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


load_local_env()
RULES = json.loads((ROOT / "rules.json").read_text(encoding="utf-8"))
MAX_REQUEST_BYTES = 25 * 1024 * 1024
MAX_TEXT_CHARS = 60_000


@dataclass
class Record:
    id: str
    source_name: str
    text: str
    extraction_warnings: tuple[str, ...] = ()


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_TEXT_CHARS]


def _read_zip_member_without_crc(data: bytes, info: zipfile.ZipInfo) -> bytes:
    """Recover a ZIP member whose payload is intact but CRC metadata is wrong."""
    header = data[info.header_offset:info.header_offset + 30]
    if len(header) != 30:
        raise zipfile.BadZipFile("DOCX成员头不完整")
    signature, *values = struct.unpack("<IHHHHHIIIHH", header)
    if signature != 0x04034B50:
        raise zipfile.BadZipFile("DOCX成员头无效")
    filename_length, extra_length = values[-2:]
    payload_start = info.header_offset + 30 + filename_length + extra_length
    compressed = data[payload_start:payload_start + info.compress_size]
    if len(compressed) != info.compress_size:
        raise zipfile.BadZipFile("DOCX成员数据不完整")
    if info.compress_type == zipfile.ZIP_STORED:
        payload = compressed
    elif info.compress_type == zipfile.ZIP_DEFLATED:
        payload = zlib.decompress(compressed, -zlib.MAX_WBITS)
    else:
        raise zipfile.BadZipFile(f"不支持修复的DOCX压缩方式：{info.compress_type}")
    return payload


def repair_docx_crc(data: bytes) -> bytes:
    """Rebuild a DOCX when one or more ZIP entries only have a bad CRC value."""
    if not zipfile.is_zipfile(io.BytesIO(data)):
        raise zipfile.BadZipFile("文件不是有效的DOCX压缩包")
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(data), "r") as source, zipfile.ZipFile(output, "w") as target:
        for info in source.infolist():
            try:
                payload = source.read(info)
            except zipfile.BadZipFile as exc:
                if "Bad CRC-32" not in str(exc):
                    raise
                payload = _read_zip_member_without_crc(data, info)
            target.writestr(info, payload)
    repaired = output.getvalue()
    with zipfile.ZipFile(io.BytesIO(repaired), "r") as verified:
        if verified.testzip() is not None:
            raise zipfile.BadZipFile("DOCX自动修复后仍未通过完整性检查")
    return repaired


def _extract_docx_archive(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    try:
        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    value = cell.text.strip().replace("\n", " / ")
                    if value and value not in cells:
                        cells.append(value)
                if cells:
                    blocks.append(" | ".join(cells))
    except InvalidSpanError:
        return _extract_docx_xml(data)
    text = clean_text("\n".join(blocks))
    if not text:
        raise ValueError("DOCX中没有读取到可审核文字")
    return text


def _extract_docx_xml(data: bytes) -> str:
    """Read paragraphs and table cells without relying on python-docx's table grid."""
    with zipfile.ZipFile(io.BytesIO(data), "r") as archive:
        document_xml = archive.read("word/document.xml")
    parser = etree.XMLParser(recover=False, huge_tree=False, resolve_entities=False)
    root = etree.fromstring(document_xml, parser)
    body_nodes = root.xpath('//*[local-name()="body"]')
    if not body_nodes:
        raise ValueError("DOCX中没有找到正文结构")
    blocks: list[str] = []
    for block in body_nodes[0]:
        block_name = etree.QName(block).localname
        if block_name == "p":
            value = "".join(block.xpath('.//*[local-name()="t"]/text()')).strip()
            if value:
                blocks.append(value)
        elif block_name == "tbl":
            for row in block.xpath('./*[local-name()="tr"]'):
                cells: list[str] = []
                for cell in row.xpath('./*[local-name()="tc"]'):
                    paragraphs = [
                        "".join(paragraph.xpath('.//*[local-name()="t"]/text()')).strip()
                        for paragraph in cell.xpath('./*[local-name()="p"]')
                    ]
                    value = " / ".join(part for part in paragraphs if part)
                    if value and value not in cells:
                        cells.append(value)
                if cells:
                    blocks.append(" | ".join(cells))
    text = clean_text("\n".join(blocks))
    if not text:
        raise ValueError("DOCX中没有读取到可审核文字")
    return text


def extract_docx_with_status(data: bytes) -> tuple[str, tuple[str, ...]]:
    try:
        return _extract_docx_archive(data), ()
    except (zipfile.BadZipFile, etree.XMLSyntaxError, KeyError):
        try:
            return _extract_docx_archive(repair_docx_crc(data)), ("DOCX压缩包校验异常，系统已完整修复后读取",)
        except (zipfile.BadZipFile, ValueError, KeyError, zlib.error, etree.XMLSyntaxError):
            office_repaired = repair_docx_with_office(data)
            if office_repaired is not None:
                try:
                    return _extract_docx_archive(office_repaired), ("DOCX结构异常，系统已使用文档转换器修复后读取",)
                except (zipfile.BadZipFile, ValueError, KeyError, etree.XMLSyntaxError):
                    pass
            try:
                return recover_corrupt_docx_text(data), ("DOCX正文结构损坏，仅恢复到部分可读内容，材料完整性无法确认",)
            except (zipfile.BadZipFile, ValueError, KeyError, zlib.error, etree.XMLSyntaxError) as exc:
                raise ValueError("DOCX正文结构损坏且无法恢复，请在Word中打开后使用“打开并修复”，再另存为新文件") from exc


def extract_docx(data: bytes) -> str:
    return extract_docx_with_status(data)[0]


def find_office_converter() -> str | None:
    configured = os.getenv("LIBREOFFICE_PATH", "").strip()
    if configured and Path(configured).is_file():
        return configured
    for command in ("libreoffice", "soffice"):
        found = shutil.which(command)
        if found:
            return found
    return None


def _run_office_conversion(data: bytes, input_suffix: str) -> bytes | None:
    converter = find_office_converter()
    if not converter:
        return None
    with tempfile.TemporaryDirectory(prefix="office-review-") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        input_dir = temp_dir / "input"
        output_dir = temp_dir / "output"
        input_dir.mkdir()
        output_dir.mkdir()
        source_path = input_dir / f"uploaded{input_suffix}"
        output_path = output_dir / "uploaded.docx"
        profile_uri = (temp_dir / "lo-profile").resolve().as_uri()
        source_path.write_bytes(data)
        completed = subprocess.run(
            [
                converter,
                "--headless",
                "--nologo",
                "--nodefault",
                "--nofirststartwizard",
                f"-env:UserInstallation={profile_uri}",
                "--convert-to",
                "docx",
                "--outdir",
                str(output_dir),
                str(source_path),
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            timeout=90,
            check=False,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        if completed.returncode != 0 or not output_path.exists():
            return None
        return output_path.read_bytes()


def repair_docx_with_office(data: bytes) -> bytes | None:
    try:
        return _run_office_conversion(data, ".docx")
    except (OSError, subprocess.TimeoutExpired):
        return None


def recover_corrupt_docx_text(data: bytes) -> str:
    """Best-effort text recovery for a DOCX whose document.xml is damaged."""
    with zipfile.ZipFile(io.BytesIO(data), "r") as archive:
        info = archive.getinfo("word/document.xml")
        try:
            document_xml = archive.read(info)
        except zipfile.BadZipFile:
            document_xml = _read_zip_member_without_crc(data, info)
    safe_xml = document_xml.decode("utf-8", errors="replace").encode("utf-8")
    parser = etree.XMLParser(recover=True, huge_tree=False, resolve_entities=False)
    root = etree.fromstring(safe_xml, parser)
    text_nodes = [str(value).strip() for value in root.xpath('//*[local-name()="t"]/text()')]
    text = clean_text("\n".join(value for value in text_nodes if value))
    if len(text) < 200:
        raise ValueError("损坏DOCX中可恢复的文字不足")
    return text


def extract_doc_with_oxide(data: bytes) -> str:
    if oxide_extract_text is None and oxide_to_markdown is None:
        raise RuntimeError("Office Oxide未安装")
    with tempfile.TemporaryDirectory(prefix="oxide-review-") as temp_dir_name:
        source_path = Path(temp_dir_name) / "uploaded.doc"
        source_path.write_bytes(data)
        candidates: list[str] = []
        for extractor in (oxide_to_markdown, oxide_extract_text):
            if extractor is None:
                continue
            try:
                value = clean_text(str(extractor(str(source_path)) or ""))
                if value:
                    candidates.append(value)
            except Exception:
                continue
    if not candidates:
        raise ValueError("Office Oxide未读取到正文")
    text = max(candidates, key=lambda value: (sum(key in value for key in ("研究内容", "预期成果", "创新点", "审批意见")), len(value)))
    if len(text) < 200:
        raise ValueError("Office Oxide提取的可审核文字不足")
    return text


def extract_doc_with_status(data: bytes) -> tuple[str, tuple[str, ...]]:
    try:
        return extract_doc_with_oxide(data), ("原始DOC由Office Oxide直接读取，未改变文件格式",)
    except (RuntimeError, ValueError, OSError):
        pass
    if not find_office_converter():
        raise ValueError("服务器未能直接读取旧版DOC，且尚未安装LibreOffice备用组件")
    try:
        converted = _run_office_conversion(data, ".doc")
    except subprocess.TimeoutExpired as exc:
        raise ValueError("旧版DOC转换超时，请确认文件未加密且可以正常打开") from exc
    if converted is None:
        raise ValueError("旧版DOC无法转换，请确认文件未加密、未损坏且不是仅修改了扩展名")
    text, warnings = extract_docx_with_status(converted)
    return text, ("Office Oxide未能直接读取，系统已使用LibreOffice临时转换后审核",) + warnings


def extract_doc(data: bytes) -> str:
    return extract_doc_with_status(data)[0]


def extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return clean_text("\n".join(page.extract_text() or "" for page in reader.pages))


def records_from_file(item: dict[str, Any]) -> list[Record]:
    name = str(item.get("name") or "uploaded")
    raw = base64.b64decode(item.get("base64") or "", validate=True)
    suffix = Path(name).suffix.lower()
    stem = Path(name).stem

    if suffix == ".docx":
        if raw.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
            text, warnings = extract_doc_with_status(raw)
            return [Record(stem, name, text, warnings)]
        text, warnings = extract_docx_with_status(raw)
        return [Record(stem, name, text, warnings)]
    if suffix == ".pdf":
        return [Record(stem, name, extract_pdf(raw))]
    if suffix == ".doc":
        text, warnings = extract_doc_with_status(raw)
        return [Record(stem, name, text, warnings)]
    if suffix in {".txt", ".md"}:
        return [Record(stem, name, clean_text(raw.decode("utf-8-sig", errors="replace")))]
    if suffix == ".json":
        value = json.loads(raw.decode("utf-8-sig"))
        rows = value if isinstance(value, list) else value.get("data", [value]) if isinstance(value, dict) else [value]
        return [
            Record(str(row.get("id", f"{stem}_{idx:03d}")) if isinstance(row, dict) else f"{stem}_{idx:03d}", name,
                   clean_text(json.dumps(row, ensure_ascii=False)))
            for idx, row in enumerate(rows, 1)
        ]
    if suffix == ".csv":
        decoded = raw.decode("utf-8-sig", errors="replace")
        rows = list(csv.DictReader(io.StringIO(decoded)))
        return [Record(str(row.get("id") or f"{stem}_{idx:03d}"), name, clean_text(json.dumps(row, ensure_ascii=False)))
                for idx, row in enumerate(rows, 1)]
    raise ValueError(f"暂不支持 {suffix or '无扩展名'} 文件：{name}")


def snippet(text: str, pattern: str, width: int = 95) -> str | None:
    match = re.search(pattern, text, re.I)
    if not match:
        return None
    start = max(0, match.start() - 20)
    end = min(len(text), match.end() + width)
    return text[start:end].replace("\n", " ").strip()


def deterministic_checks(
    text: str,
    dataset_type: str,
    extraction_warnings: tuple[str, ...] = (),
) -> list[dict[str, str]]:
    checks: list[dict[str, str]] = []
    for warning in extraction_warnings:
        if "仅恢复到部分" in warning or "完整性无法确认" in warning:
            checks.append({"rule_id": "R-C02", "message": "文档损坏导致内容不完整", "evidence": warning})
    patterns = [
        ("R-C01", r"填写说明|删除本提示|待补充|请在此处填写", "存在未清理的模板提示或占位内容"),
        ("R-C04", r"百分之|准确率\s*\|\s*\|", "考核指标可能缺失、格式异常或不可量化"),
        ("R-C07", r"\n1\s*\|?\s*[^\n]+\n1\s*\|", "人员或事项序号重复"),
        ("R-C08", r"研究报告\s*\|\s*初稿", "验收成果状态仍为初稿"),
        ("R-C09", r"不同意立项|不予立项|退回修改", "审批意见包含明确否决信息"),
    ]
    for rule_id, pattern, message in patterns:
        evidence = snippet(text, pattern)
        if evidence:
            checks.append({"rule_id": rule_id, "message": message, "evidence": evidence})

    if dataset_type == "立项申请书":
        approval_headings = [
            ("申请部门/单位意见", r"申请(?:部门/单位|部门|单位)意见"),
            ("科技管理部门意见", r"(?:直属单位|申请单位)科技管理部门意见"),
        ]
        blank_approval_sections: list[str] = []
        for index, (label, heading_pattern) in enumerate(approval_headings):
            next_heading_pattern = approval_headings[index + 1][1] if index + 1 < len(approval_headings) else None
            end_pattern = next_heading_pattern if next_heading_pattern else r"注[：:]|廉洁及科研诚信承诺书|$"
            section = re.search(
                rf"{heading_pattern}[：:]?(.*?)(?={end_pattern})",
                text,
                re.S,
            )
            if section and not re.search(r"同意|不同意|不予|退回|经审核|批准|建议", section.group(1)):
                blank_approval_sections.append(label)
        if len(blank_approval_sections) == len(approval_headings):
            checks.append({
                "rule_id": "R-C09",
                "message": "两级审批意见均未填写明确结论",
                "evidence": "“申请部门/单位意见”和“直属单位科技管理部门意见”均只有公章或日期占位，未填写明确审批意见",
            })

    if dataset_type == "计划任务书":
        if "项目执行期限" not in text and "项目起止时间" not in text:
            checks.append({"rule_id": "R-C05", "message": "缺少项目执行期限", "evidence": "未检索到项目执行期限或项目起止时间"})
        if "考核指标" not in text:
            checks.append({"rule_id": "R-C04", "message": "缺少可验收考核指标", "evidence": "未检索到考核指标章节"})

        unit_values = re.findall(r"项目承担单位\s*\|\s*([^\n|]+)", text)
        normalized_units = {value.strip() for value in unit_values}
        if len(normalized_units) > 1:
            checks.append({"rule_id": "R-C03", "message": "项目承担单位前后不一致", "evidence": "；".join(sorted(normalized_units))})
        elif len(normalized_units) == 1:
            stated_unit = next(iter(normalized_units))
            full_names = set(re.findall(r"([\u4e00-\u9fff]{4,30}(?:有限责任公司|供电局|研究院|管理所))", text))
            longer_matches = sorted(name for name in full_names if name.endswith(stated_unit) and name != stated_unit)
            if longer_matches:
                checks.append({"rule_id": "R-C03", "message": "承担单位简称与其他章节全称不一致", "evidence": f"项目承担单位填写“{stated_unit}”，其他章节出现“{longer_matches[0]}”"})

        profile_people = {
            name: age for name, age in re.findall(r"\n([\u4e00-\u9fff]{2,4})\s*\|\s*(?:男|女)\s*\|\s*(\d{1,2})\s*\|", text)
        }
        team_people = {
            name: age for name, age in re.findall(r"\n\d+\s*\|\s*([\u4e00-\u9fff]{2,4})\s*\|\s*(\d{1,2})\s*\|", text)
        }
        for name, age in profile_people.items():
            if name in team_people and team_people[name] != age:
                checks.append({"rule_id": "R-C03", "message": "负责人年龄前后不一致", "evidence": f"{name}在基本信息中为{age}岁，在团队表中为{team_people[name]}岁"})

        team_section = re.search(r"序号\s*\|\s*姓名.*?(?=序号\s*\|\s*考核指标名称)", text, re.S)
        if team_section:
            ids = re.findall(r"\n(\d+)\s*\|", team_section.group(0))
            duplicates = sorted({item for item in ids if ids.count(item) > 1})
            if duplicates:
                checks.append({"rule_id": "R-C07", "message": "团队人员序号重复", "evidence": f"重复序号：{', '.join(duplicates)}"})
            if len(ids) > 20:
                checks.append({"rule_id": "R-C07", "message": "团队人员数量异常", "evidence": f"团队表共列出{len(ids)}行人员"})

        if re.search(r"外部专家.*?(?:茂名|佛山|广州|珠海|东莞)供电局", text, re.S):
            evidence = snippet(text, r"外部专家.*?(?:茂名|佛山|广州|珠海|东莞)供电局")
            checks.append({"rule_id": "R-C07", "message": "存在未在承担单位中说明的外部团队成员", "evidence": evidence or "发现外部专家单位"})

        duration = re.search(r"项目执行期限\s*\|\s*(\d{4})-(\d{2})\s*至\s*(\d{4})-(\d{2})", text)
        schedule_section = re.search(r"序号\s*\|\s*时间段.*?(?=经费类型\s*\|)", text, re.S)
        if duration and schedule_section:
            months = (int(duration.group(3)) - int(duration.group(1))) * 12 + int(duration.group(4)) - int(duration.group(2)) + 1
            stages = re.findall(r"\n\d+\s*\|\s*\d{4}-\d{2}\s*至", schedule_section.group(0))
            if months > 12 and len(stages) < 3:
                checks.append({"rule_id": "R-C05", "message": "进度安排未覆盖完整执行期", "evidence": f"项目执行期约{months}个月，但仅列出{len(stages)}个进度阶段"})

        if duration:
            months = (int(duration.group(3)) - int(duration.group(1))) * 12 + int(duration.group(4)) - int(duration.group(2)) + 1
            if months > 30:
                checks.append({"rule_id": "R-C05", "message": "项目执行期超过常见30个月范围，需复核", "evidence": duration.group(0)})
    else:
        for keyword, label in [("研究内容", "研究内容"), ("技术关键点及创新点", "创新点"), ("应用前景", "应用前景")]:
            if keyword not in text:
                checks.append({"rule_id": "R-A01", "message": f"缺少{label}", "evidence": f"未检索到“{keyword}”章节"})
    return checks


def select_rules(dataset_type: str, intent: str) -> list[dict[str, Any]]:
    words = set(re.findall(r"[\u4e00-\u9fff]{2,}|[A-Za-z]{3,}", intent))
    ranked = []
    for rule in RULES:
        if dataset_type not in rule["dataset_types"] and "通用" not in rule["dataset_types"]:
            continue
        score = sum(2 for word in words if word in rule["name"] or word in rule["description"])
        score += 1 if rule["severity"] == "blocking" else 0
        ranked.append((score, rule))
    ranked.sort(key=lambda item: (-item[0], item[1]["rule_id"]))
    return [item[1] for item in ranked[:10]]


def parse_json_response(content: str) -> dict[str, Any]:
    content = content.strip()
    content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content, flags=re.I)
    start, end = content.find("{"), content.rfind("}")
    if start < 0 or end <= start:
        raise ValueError("模型没有返回 JSON 对象")
    return json.loads(content[start:end + 1])


def call_llm(
    record: Record,
    dataset_type: str,
    intent: str,
    checks: list[dict[str, str]],
    review_note: str | None = None,
) -> dict[str, Any]:
    api_key = os.getenv("LLM_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("LLM_API_KEY 未配置")
    base_url = os.getenv("LLM_BASE_URL", "https://api.deepseek.com").rstrip("/")
    model = os.getenv("LLM_MODEL", "deepseek-chat")
    candidate_rules = select_rules(dataset_type, intent)
    system = """你是业务文档合规判断智能体。必须先分析intent，再提取事实，分别寻找否决证据和通过条件，匹配规则，最后输出硬标签。不能根据项目题目热门程度猜测。evidence必须是文档中的短原文或明确的缺失事实，不得编造。仅输出一个JSON对象，不输出Markdown。"""
    prompt = {
        "task": "根据intent审核文档并返回固定结构结果",
        "dataset_type": dataset_type,
        "intent": intent,
        "record_id": record.id,
        "candidate_rules": candidate_rules,
        "deterministic_checks": checks,
        "decision_policy": [
            "按顺序完成：提取关键事实、检查硬性缺陷、评估问题方案匹配、评估创新实质、核对一致性，最后裁决",
            "先检查blocking规则；命中且证据充分通常判不通过",
            "未发现否决证据不等于自动通过；通过必须有充分正面证据",
            "信息缺失影响核心判断时判不通过，不要猜测",
            "matched_rules只放真正用于裁决且有证据的规则",
        ],
        "calibration_guardrails": [
            "签字、公章或日期空白不能单独否决；但审批意见栏只有公章或日期占位，没有同意、不同意、退回、建议等意见内容时，属于审批意见本身空白；立项申请书两级意见均空白应按R-C09审查",
            "某标题附近没有正文不等于内容缺失；必须检查全文其他表格或章节是否已提供对应信息",
            "预算比较必须确认统计口径相同；费用性明细不能直接与包含资本性的总经费比较",
            "资本性经费未在费用性年度栏展示，不得自动推断预算不闭合",
            "非关键格式瑕疵不能替代与当前intent有关的实质否决证据",
        ],
        "few_shot_principles": [
            "若材料只有现成部件组合、通用功能罗列和泛化效益，没有可区分的关键机制或验证指标，应按创新实质不足审查R-A03",
            "若方案给出明确业务痛点、可区分技术机制、可验证指标且跨章节一致，不能仅因题目普通或签章日期空白否决",
        ],
        "required_output": {
            "label": "通过 或 不通过",
            "matched_rules": [{"rule_id": "规则编号", "rule_name": "规则名", "evidence": "文档原文证据"}],
            "reason": "不超过120字",
            "confidence": "0到1之间的小数",
        },
        "document": record.text,
    }
    if review_note:
        prompt["independent_review"] = review_note
    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": json.dumps(prompt, ensure_ascii=False)},
    ]
    last_error: Exception | None = None
    for attempt in range(2):
        if attempt:
            messages.append({"role": "user", "content": "上一次返回无法解析。请严格只返回一个有效JSON对象，不要输出解释、代码围栏或思考过程。"})
        body = json.dumps({
            "model": model,
            "temperature": 0,
            "max_tokens": 1400,
            "response_format": {"type": "json_object"},
            "messages": messages,
        }, ensure_ascii=False).encode("utf-8")
        request = urllib.request.Request(
            f"{base_url}/chat/completions",
            data=body,
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(request, timeout=90) as response:
                payload = json.loads(response.read().decode("utf-8"))
            content = payload["choices"][0]["message"].get("content") or ""
            return parse_json_response(content)
        except urllib.error.HTTPError as exc:
            messages_by_status = {
                400: "DeepSeek拒绝了请求参数",
                401: "DeepSeek API Key无效或未生效",
                402: "DeepSeek账户余额不足",
                429: "DeepSeek请求过于频繁，请稍后重试",
            }
            raise RuntimeError(messages_by_status.get(exc.code, f"DeepSeek接口返回HTTP {exc.code}")) from exc
        except urllib.error.URLError as exc:
            raise RuntimeError("无法连接DeepSeek，请检查网络、代理或防火墙") from exc
        except (ValueError, KeyError, json.JSONDecodeError) as exc:
            last_error = exc
    raise RuntimeError(f"模型连续两次未返回有效JSON：{last_error}")


def heuristic_result(record: Record, dataset_type: str, intent: str, checks: list[dict[str, str]]) -> dict[str, Any]:
    blocking = {rule["rule_id"]: rule for rule in RULES if rule["severity"] == "blocking"}
    matched = []
    seen = set()
    for check in checks:
        rule = blocking.get(check["rule_id"])
        signature = (check["rule_id"], check["evidence"])
        if rule and signature not in seen:
            seen.add(signature)
            matched.append({"rule_id": rule["rule_id"], "rule_name": rule["name"], "evidence": check["evidence"]})
    label = "不通过" if matched else "通过"
    reason = "发现影响材料有效性或可验收性的明确问题。" if matched else "本地预检查未发现明确否决项；正式结果仍建议由模型复核。"
    return {"label": label, "matched_rules": matched, "reason": reason, "confidence": 0.82 if matched else 0.56}


def normalize_result(raw: dict[str, Any], record: Record, dataset_type: str, intent: str, mode: str) -> dict[str, Any]:
    label = raw.get("label") if raw.get("label") in {"通过", "不通过"} else "不通过"
    rules = raw.get("matched_rules") if isinstance(raw.get("matched_rules"), list) else []
    normalized_rules = []
    for item in rules[:8]:
        if isinstance(item, dict):
            normalized_rules.append({
                "rule_id": str(item.get("rule_id", "R-UNKNOWN")),
                "rule_name": str(item.get("rule_name", "未命名规则")),
                "evidence": str(item.get("evidence", "未提供证据"))[:300],
            })
    try:
        confidence = max(0.0, min(1.0, float(raw.get("confidence", 0.5))))
    except (TypeError, ValueError):
        confidence = 0.5
    result = {
        "id": record.id,
        "source_name": record.source_name,
        "dataset_type": dataset_type,
        "intent": intent,
        "label": label,
        "matched_rules": normalized_rules,
        "reason": str(raw.get("reason", "未提供判断说明"))[:300],
        "confidence": confidence,
        "mode": mode,
    }
    if record.extraction_warnings:
        result["warning"] = "；".join(record.extraction_warnings)
    return result


def judge_one(record: Record, dataset_type: str, intent: str) -> dict[str, Any]:
    checks = deterministic_checks(record.text, dataset_type, record.extraction_warnings)
    if os.getenv("LLM_API_KEY", "").strip():
        try:
            raw = call_llm(record, dataset_type, intent, checks)
            approval_checks = [check for check in checks if check["rule_id"] == "R-C09"]
            if approval_checks and raw.get("label") == "通过":
                raw = {
                    "label": "不通过",
                    "matched_rules": [
                        {
                            "rule_id": "R-C09",
                            "rule_name": "审批意见有效性",
                            "evidence": check["evidence"],
                        }
                        for check in approval_checks
                    ],
                    "reason": "立项申请书的审批意见缺少明确结论，无法确认已满足立项审批要求。",
                    "confidence": 0.96,
                }
            if raw.get("label") == "通过" and not checks and dataset_type == "立项申请书":
                first_result = json.dumps(raw, ensure_ascii=False)
                critical_result = call_llm(
                    record,
                    dataset_type,
                    intent,
                    checks,
                    review_note=(
                        "这是一次防漏判的独立反向复核。不要因为栏目齐全就自动通过。"
                        "重点检查问题与方案是否真正匹配、创新点是否只是现成设备或模块的简单拼装、"
                        "技术机制和验证指标是否足以支撑声称效果，以及跨章节事实是否一致。"
                        "只有找到文档中的明确事实证据才能判不通过；不要根据文件名或已知标签猜测。"
                        f"主审核结果：{first_result}"
                    ),
                )
                if critical_result.get("label") == "不通过":
                    raw = call_llm(
                        record,
                        dataset_type,
                        intent,
                        checks,
                        review_note=(
                            "主审核与反向复核结论冲突，请作为独立裁决智能体重新核对原文证据。"
                            "不能用签字、公章或日期空白单独否决；但两级审批意见均无明确结论时属于R-C09。也不能因栏目齐全自动通过。"
                            "若存在问题方案不匹配、创新实质不足或关键事实矛盾，必须引用对应短原文并匹配规则。"
                            f"主审核：{first_result}；反向复核：{json.dumps(critical_result, ensure_ascii=False)}"
                        ),
                    )
                else:
                    raw = critical_result
            elif raw.get("label") == "不通过" and not checks:
                first_result = json.dumps(raw, ensure_ascii=False)
                raw = call_llm(
                    record,
                    dataset_type,
                    intent,
                    checks,
                    review_note=(
                        "这是一次反误杀独立复核。程序未发现高置信度结构问题。"
                        "请质疑第一次结论是否把签字日期空白、跨章节内容位置或不同预算口径误当成否决证据。"
                        "如果仍有与intent直接相关、证据充分的实质问题，可以维持不通过；否则改为通过。"
                        f"第一次结果：{first_result}"
                    ),
                )
            return normalize_result(raw, record, dataset_type, intent, "llm")
        except Exception as exc:
            fallback = heuristic_result(record, dataset_type, intent, checks)
            result = normalize_result(fallback, record, dataset_type, intent, "fallback")
            model_warning = f"模型调用失败，已使用本地预检查：{str(exc)[:160]}"
            result["warning"] = "；".join(filter(None, [result.get("warning"), model_warning]))
            return result
    return normalize_result(heuristic_result(record, dataset_type, intent, checks), record, dataset_type, intent, "demo")


class AppHandler(SimpleHTTPRequestHandler):
    def translate_path(self, path: str) -> str:
        clean = path.split("?", 1)[0].split("#", 1)[0]
        if clean == "/":
            clean = "/index.html"
        return str(WEB_ROOT / clean.lstrip("/"))

    def send_json(self, value: Any, status: int = 200) -> None:
        data = json.dumps(value, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def do_GET(self) -> None:
        if self.path == "/api/health":
            self.send_json({
                "ok": True,
                "mode": "llm" if os.getenv("LLM_API_KEY") else "demo",
                "provider": "DeepSeek" if "deepseek" in os.getenv("LLM_BASE_URL", "https://api.deepseek.com").lower() else "OpenAI-compatible",
                "model": os.getenv("LLM_MODEL", "deepseek-chat"),
            })
            return
        super().do_GET()

    def do_POST(self) -> None:
        if self.path != "/api/analyze":
            self.send_json({"error": "接口不存在"}, HTTPStatus.NOT_FOUND)
            return
        try:
            length = int(self.headers.get("Content-Length", "0"))
            if length <= 0 or length > MAX_REQUEST_BYTES:
                raise ValueError("请求为空或超过25MB限制")
            payload = json.loads(self.rfile.read(length).decode("utf-8"))
            dataset_type = str(payload.get("dataset_type", ""))
            intent = str(payload.get("intent", "")).strip()
            files = payload.get("files", [])
            if dataset_type not in {"计划任务书", "立项申请书"}:
                raise ValueError("请选择正确的数据集类型")
            if not intent:
                raise ValueError("请输入判断 intent")
            if not isinstance(files, list) or not files:
                raise ValueError("请至少上传一个文件")
            work_items: list[Record | dict[str, Any]] = []
            for item in files:
                try:
                    work_items.extend(records_from_file(item))
                except Exception as exc:
                    name = str(item.get("name") or "uploaded") if isinstance(item, dict) else "uploaded"
                    work_items.append({
                        "id": Path(name).stem,
                        "source_name": name,
                        "dataset_type": dataset_type,
                        "intent": intent,
                        "label": "处理失败",
                        "matched_rules": [],
                        "reason": f"文件读取失败：{str(exc)[:220] or type(exc).__name__}",
                        "confidence": 0.0,
                        "mode": "error",
                    })
            if not work_items:
                raise ValueError("上传文件中没有可判断的数据")
            results: list[dict[str, Any] | None] = [None] * len(work_items)
            records = [(idx, item) for idx, item in enumerate(work_items) if isinstance(item, Record)]
            for idx, item in enumerate(work_items):
                if isinstance(item, dict):
                    results[idx] = item
            with ThreadPoolExecutor(max_workers=max(1, min(3, len(records)))) as pool:
                futures = {pool.submit(judge_one, record, dataset_type, intent): idx for idx, record in records}
                for future in as_completed(futures):
                    results[futures[future]] = future.result()
            self.send_json({
                "results": results,
                "count": len(results),
                "mode": "llm" if os.getenv("LLM_API_KEY") else "demo",
            })
        except (ValueError, json.JSONDecodeError, base64.binascii.Error) as exc:
            self.send_json({"error": str(exc)}, HTTPStatus.BAD_REQUEST)
        except Exception as exc:
            self.send_json({"error": f"系统处理失败：{type(exc).__name__}"}, HTTPStatus.INTERNAL_SERVER_ERROR)


def main() -> None:
    host = os.getenv("APP_HOST", "127.0.0.1")
    port = int(os.getenv("PORT", "8000"))
    server = ThreadingHTTPServer((host, port), AppHandler)
    mode = "LLM" if os.getenv("LLM_API_KEY") else "演示"
    print(f"业务文档审核系统已启动：http://{host}:{port}（{mode}模式）", flush=True)
    server.serve_forever()


if __name__ == "__main__":
    main()
